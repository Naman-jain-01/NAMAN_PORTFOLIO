import os
import re
from docx import Document
import argparse

def extract_placeholders_from_docx(filepath):
    placeholders = set()
    pattern = re.compile(r'\{\{\s*(.+?)\s*\}\}')

    try:
        document = Document(filepath)
        for para in document.paragraphs:
            found = pattern.findall(para.text)
            placeholders.update(found)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    found = pattern.findall(cell.text)
                    placeholders.update(found)
    except Exception as e:
        print(f"Error processing file {filepath}: {e}")
    return placeholders

def find_placeholders_in_directory(directory, output_path):
    all_placeholders = set()
    for filename in os.listdir(directory):
        if filename.endswith('.docx'):
            filepath = os.path.join(directory, filename)
            placeholders = extract_placeholders_from_docx(filepath)
            all_placeholders.update(placeholders)
    
    # Save all placeholders to a text file
    if os.path.exists(output_path):
        os.remove(output_path)
    with open(output_path, 'w') as file:
        file.write(', '.join(all_placeholders))

def main():
    parser = argparse.ArgumentParser(description="Extract placeholders from all DOCX files in a directory and save them to a specified output path.")
    parser.add_argument("--directory", type=str, help="Directory containing DOCX files")
    parser.add_argument("--output_path", type=str, help="Full path to the output file")

    args = parser.parse_args()

    find_placeholders_in_directory(args.directory, args.output_path)
    print(f"Placeholders have been written to {args.output_path}")

if __name__ == "__main__":
    main()


import sys
sys.path.append("/workspaces/NAMAN_PORTFOLIO/server/BKL/lib/python3.10/site-packages")

import pandas as pd
import re
import io
import os
from docx import Document
import zipfile
import argparse


def replace_paragraph_text(paragraph, replacements):
    for key, value in replacements.items():
        pattern = f'{{{{ {key} }}}}'  # Find placeholders like '{{ key }}'
        if pattern in paragraph.text:
            new_text = paragraph.text
            # Replace each placeholder with its corresponding value
            new_text = re.sub(re.escape(pattern), value, new_text)
            paragraph.clear()
            # Split the text around the value to isolate it for formatting
            parts = re.split(f'({re.escape(value)})', new_text)
            for part in parts:
                # Add each part back to the paragraph
                run = paragraph.add_run(part)
            

def replace_cell_text(cell, replacements):
    for paragraph in cell.paragraphs:
        replace_paragraph_text(paragraph, replacements)

def process_templates_from_folders(template_folder, excel_folder, output_format, output_folder):
    excel_files = [f for f in os.listdir(excel_folder) if f.endswith('.xlsx')]
    if not excel_files:
        raise FileNotFoundError("No Excel files found in the specified folder.")
    excel_path = os.path.join(excel_folder, excel_files[0])

    data = pd.read_excel(excel_path)
    data.columns = data.columns.str.strip()

    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    for index, row in data.iterrows():
        zip_path = os.path.join(output_folder, f"Row_WISE_{index + 1}_files.zip")

        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for template_file in os.listdir(template_folder):
                if template_file.endswith('.docx'):
                    template_name = os.path.splitext(template_file)[0]
                    template_path = os.path.join(template_folder, template_file)
                    new_doc = Document(template_path)
                    # Use row.to_dict() to ensure the row data is correctly formatted as a dict
                    replacements = {column: str(row[column]) for column in data.columns if column in row}

                    for paragraph in new_doc.paragraphs:
                        replace_paragraph_text(paragraph, replacements)
                    for table in new_doc.tables:
                        for table_row in table.rows:
                            for cell in table_row.cells:
                                replace_cell_text(cell, replacements)

                    docx_stream = io.BytesIO()
                    new_doc.save(docx_stream)
                    docx_stream.seek(0)

                    file_name = f"{template_name}_{index + 1}"

                    if output_format.lower() == 'pdf':
                        zipf.writestr(f"{file_name}.docx", docx_stream.read())
                    else:
                        zipf.writestr(f"{file_name}.docx", docx_stream.read())

    excel_files = [f for f in os.listdir(excel_folder) if f.endswith('.xlsx')]
    if not excel_files:
        raise FileNotFoundError("No Excel files found in the specified folder.")
    excel_path = os.path.join(excel_folder, excel_files[0])

    data = pd.read_excel(excel_path)
    data.columns = data.columns.str.strip()

    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    for template_file in os.listdir(template_folder):
        if template_file.endswith('.docx'):
            template_name = os.path.splitext(template_file)[0]
            zip_path = os.path.join(output_folder, f"{template_name}_files.zip")

            with zipfile.ZipFile(zip_path, 'w') as zipf:
                for index, row in data.iterrows():
                    template_path = os.path.join(template_folder, template_file)
                    new_doc = Document(template_path)
                    replacements = {column: str(row[column]) for column in data.columns}

                    for paragraph in new_doc.paragraphs:
                        replace_paragraph_text(paragraph, replacements)
                    for table in new_doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                replace_cell_text(cell, replacements)

                    docx_stream = io.BytesIO()
                    new_doc.save(docx_stream)
                    docx_stream.seek(0)

                    file_name = f"{template_name}_{index + 1}"
                    if output_format.lower() == 'pdf':
                        zipf.writestr(f"{file_name}.docx", docx_stream.read())
                    else:
                        zipf.writestr(f"{file_name}.docx", docx_stream.read())

# Argument parsing
if __name__ == '__main__':
    parser = argparse.ArgumentParser(description="Generate documents based on templates and Excel data.")
    parser.add_argument('--template_folder', required=True, help="Folder containing the DOCX templates")
    parser.add_argument('--excel_folder', required=True, help="Folder containing the Excel file")
    parser.add_argument('--output_format', required=True, choices=['pdf', 'docx'], help="Output format: pdf or docx")
    parser.add_argument('--output_folder', required=True, help="Folder where the ZIP files will be stored")

    args = parser.parse_args()

process_templates_from_folders(

        template_folder=args.template_folder,
        excel_folder=args.excel_folder,
        output_format=args.output_format,
        output_folder=args.output_folder
)